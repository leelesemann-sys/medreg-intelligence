import streamlit as st
import os
import io
import zipfile
import requests
from datetime import datetime
from dotenv import load_dotenv
import cohere
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- IMPORTS ---
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
try:
    from langchain_chroma import Chroma
except ImportError:
    from langchain_community.vectorstores import Chroma

from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import HumanMessage, AIMessage
from chunking import chunk_documents

# --- KONFIGURATION ---
load_dotenv()
DB_DIR = "./db_storage"
DB_ZIP_URL = "https://github.com/leelesemann-sys/medreg-intelligence/releases/download/v1.0.0/db_storage.zip"
API_VERSION = "2024-08-01-preview"

# Keys: .env (lokal) oder Streamlit Secrets (Cloud)
def get_secret(key):
    """Liest einen Secret-Wert: zuerst aus Streamlit Secrets, dann aus .env/Umgebung."""
    try:
        return st.secrets[key]
    except (KeyError, FileNotFoundError):
        return os.getenv(key, "")

AZURE_API_KEY = get_secret("AZURE_OPENAI_API_KEY")
AZURE_ENDPOINT = get_secret("AZURE_OPENAI_ENDPOINT")
COHERE_API_KEY = get_secret("COHERE_API_KEY")

os.environ["AZURE_OPENAI_API_KEY"] = AZURE_API_KEY
os.environ["AZURE_OPENAI_ENDPOINT"] = AZURE_ENDPOINT
os.environ["AZURE_OPENAI_API_VERSION"] = API_VERSION

# --- COHERE RERANKER ---
reranker = cohere.ClientV2(api_key=COHERE_API_KEY) if COHERE_API_KEY else None

st.set_page_config(page_title="MedReg Intelligence", layout="wide", page_icon="⚖️")

# --- Minimales CSS: nur File-Uploader Fix ---
st.markdown("""
<style>
    [data-testid="stFileUploaderFileList"] {
        max-height: 400px;
        overflow-y: auto;
    }
    [data-testid="stFileUploaderFileList"] + div[data-testid] {
        display: none;
    }
</style>
""", unsafe_allow_html=True)

st.title("⚖️ MedReg Intelligence")

# --- LESBARE DOKUMENTNAMEN ---
FRIENDLY_NAMES = {
    # Vorindexierte Dokumente
    "CELEX_32017R0745_DE_TXT.pdf": "EU MDR (Verordnung 2017/745) – DE",
    "CELEX_32017R0745_EN_TXT.pdf": "EU MDR (Regulation 2017/745) – EN",
    "DE_MedProdGesetz_2020_German.pdf": "MPDG (Medizinprodukterecht-Durchführungsgesetz)",
    "fedlex-data-admin-ch-eli-cc-2020-552-20231101-de-pdf-a-5.pdf": "MepV (Medizinprodukteverordnung, Schweiz)",
    "UK_MDR_2002_Conformity_Assessment_English.pdf": "UK MDR 2002 – Conformity Assessment",
    "UK_MedDevReg_2002_English.pdf": "UK MDR 2002 (Medical Devices Regulations)",
    "mdcg_2021-24_en_0.pdf": "MDCG 2021-24 (Classification Guidance)",
    "Guidance_on_the_regulation_of_IVD_medical_devices_in_GB.pdf": "UK IVD Guidance (MHRA)",
    # Häufig hochgeladene Test-Dokumente
    "0718 FDA 21 CFR 820.pdf": "FDA 21 CFR 820 (Quality System Regulation)",
    "0718-fda-21-cfr-820.pdf": "FDA 21 CFR 820 (Quality System Regulation)",
    "sor-98-282.pdf": "CMDR (Canadian Medical Devices Regulations)",
    "CELEX_32017R0746_DE_TXT.pdf": "EU IVDR (Verordnung 2017/746) – DE",
    "CELEX_32017R0746_EN_TXT.pdf": "EU IVDR (Regulation 2017/746) – EN",
}

def friendly_name(raw_name):
    """Gibt einen lesbaren Namen zurück, oder den Originalnamen falls unbekannt."""
    return FRIENDLY_NAMES.get(raw_name, raw_name)

# --- ROBUSTER HTML EXPORT ---
def generate_audit_html(history, files):
    now = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; color: #333; }}
            .header {{ background: linear-gradient(135deg, #1a3a5c 0%, #2980b9 100%); color: white; padding: 30px; border-radius: 5px; margin-bottom: 30px; }}
            .header h1 {{ margin: 0; font-size: 24px; }}
            .header .brand {{ font-size: 0.7em; font-weight: 300; opacity: 0.85; margin-top: 4px; }}
            .meta {{ color: #eee; font-size: 0.9em; margin-top: 10px; }}
            .section-title {{ border-bottom: 2px solid #2980b9; color: #2980b9; padding-bottom: 5px; margin-top: 30px; }}
            table {{ border-collapse: collapse; width: 100%; margin: 20px 0; font-size: 0.9em; }}
            th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; vertical-align: top; }}
            th {{ background-color: #f2f2f2; font-weight: bold; }}
            .chat-entry {{ margin-bottom: 25px; padding: 15px; border-radius: 5px; }}
            .user {{ background-color: #f8f9fa; border-left: 5px solid #bdc3c7; }}
            .ai {{ background-color: #e3f2fd; border-left: 5px solid #2980b9; }}
            .role {{ font-weight: bold; margin-bottom: 10px; font-size: 0.8em; text-transform: uppercase; color: #7f8c8d; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>⚖️ MedReg Intelligence</h1>
            <div class="brand">Lesemann AI Solutions & Consulting</div>
            <div class="meta">Regulatory Audit Trail | Erstellt: {now}</div>
        </div>

        <h3 class="section-title">Datenbasis (Analysierte Gesetze)</h3>
        <ul>{"".join([f"<li>{f}</li>" for f in files])}</ul>

        <h3 class="section-title">Protokollierter Beratungsverlauf</h3>
    """
    for msg in history:
        role_class = "user" if msg.type == "human" else "ai"
        role_label = "Nutzerfrage" if msg.type == "human" else "KI-Analyse (MedReg Intelligence)"
        content = msg.content.replace("\n", "<br>")
        html += f"""
        <div class="chat-entry {role_class}">
            <div class="role">{role_label}</div>
            <div class="text">{content}</div>
        </div>"""

    html += "</body></html>"
    return html

# --- INITIALISIERUNG ---
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "processed_files" not in st.session_state:
    st.session_state.processed_files = []
if "indexing_done" not in st.session_state:
    st.session_state.indexing_done = False

# --- UI: SIDEBAR ---

with st.sidebar:
    st.header("MedReg Intelligence")

    # API-Check (nur Fehler anzeigen, Erfolg ist selbstverständlich)
    if not (AZURE_API_KEY and AZURE_ENDPOINT):
        st.error("⚠️ API Keys fehlen! Bitte .env oder Streamlit Secrets konfigurieren.")
        st.stop()

    if st.session_state.chat_history:
        st.divider()
        html_report = generate_audit_html(st.session_state.chat_history, st.session_state.processed_files)
        st.download_button(
            label="📄 Audit Trail herunterladen",
            data=html_report,
            file_name=f"Audit_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
            mime="text/html"
        )


# --- VORBERECHNETE WISSENSBASIS (Auto-Download) ---
def ensure_default_db():
    """Lädt die vorberechnete ChromaDB von GitHub Releases, falls lokal nicht vorhanden."""
    if os.path.exists(DB_DIR) and os.listdir(DB_DIR):
        return True
    try:
        zip_path = "db_storage.zip"
        with st.spinner("📥 Lade vorberechnete Wissensbasis herunter..."):
            resp = requests.get(DB_ZIP_URL, stream=True, timeout=60)
            resp.raise_for_status()
            with open(zip_path, "wb") as f:
                for chunk in resp.iter_content(chunk_size=8192):
                    f.write(chunk)
            with zipfile.ZipFile(zip_path, "r") as z:
                z.extractall(DB_DIR)
            os.remove(zip_path)
        return True
    except Exception as e:
        st.warning(f"⚠️ Auto-Download fehlgeschlagen: {e}")
        return False

# --- VEKTORDATENBANK ---
embeddings = AzureOpenAIEmbeddings(
    azure_deployment="text-embedding-3-small",
    openai_api_version=API_VERSION
)

# 1. Versuche vorberechnete DB zu laden (oder herunterladen)
db_ready = ensure_default_db()

if db_ready:
    vectorstore = Chroma(persist_directory=DB_DIR, embedding_function=embeddings)
    st.session_state.indexing_done = True

    try:
        content_data = vectorstore.get(include=['metadatas'])
        if content_data['metadatas']:
            unique_files = sorted(list(set([m['source_document'] for m in content_data['metadatas'] if 'source_document' in m])))
            st.session_state.processed_files = unique_files
    except Exception:
        pass
else:
    vectorstore = None

# --- SIDEBAR: Wissensbasis-Info ---
with st.sidebar:
    st.divider()

    # Dokumentenliste (immer sichtbar, lesbare Namen)
    if st.session_state.processed_files:
        st.markdown(f"**📚 Wissensbasis** ({len(st.session_state.processed_files)} Dokumente)")
        for f in st.session_state.processed_files:
            st.caption(f"• {friendly_name(f)}")
    else:
        st.info("Noch keine Dokumente geladen.")

    # Erfolgsmeldung nach Upload (bleibt bis nächster Rerun)
    if "upload_success" in st.session_state and st.session_state.upload_success:
        st.success(st.session_state.upload_success)
        st.session_state.upload_success = None

    st.divider()

    # --- Upload: Eigene Dokumente ---
    st.markdown("**📂 Eigene Dokumente hinzufügen**")
    st.caption("PDFs hochladen — sie werden automatisch analysiert und zur Wissensbasis hinzugefügt.")
    uploaded_files = st.file_uploader(
        "PDFs auswählen",
        type="pdf",
        accept_multiple_files=True,
        label_visibility="collapsed"
    )

    if uploaded_files:
        # Track welche Dateien schon verarbeitet wurden
        if "already_indexed_files" not in st.session_state:
            st.session_state.already_indexed_files = set()

        new_files = [f for f in uploaded_files if f.name not in st.session_state.already_indexed_files]

        if new_files:
            st.info(f"📄 {len(new_files)} neue Datei(en) erkannt. Verarbeitung startet...")

            with st.spinner("Schritt 1/2: Dokumente werden analysiert und in Abschnitte aufgeteilt..."):
                all_splits = chunk_documents(new_files)

            st.info(f"✂️ {len(all_splits)} Textabschnitte erstellt. Indexierung läuft...")

            prog = st.progress(0, text="Schritt 2/2: Textabschnitte werden indexiert...")
            batch_size = 50
            for i in range(0, len(all_splits), batch_size):
                batch = all_splits[i:i + batch_size]
                if vectorstore is None and i == 0:
                    vectorstore = Chroma.from_documents(documents=batch, embedding=embeddings, persist_directory=DB_DIR)
                else:
                    vectorstore.add_documents(batch)
                prog.progress(min((i + batch_size) / len(all_splits), 1.0))

            # Merke verarbeitete Dateien
            for f in new_files:
                st.session_state.already_indexed_files.add(f.name)

            st.session_state.indexing_done = True
            file_names = ", ".join([f.name for f in new_files])
            st.session_state.upload_success = f"✅ Fertig! {len(new_files)} Dokument(e) hinzugefügt: {file_names}"
            st.rerun()
        else:
            st.caption("✅ Alle hochgeladenen Dateien sind bereits in der Wissensbasis.")

# --- EXPORT-HELFER ---
def generate_answer_docx(question, answer):
    """Erzeugt ein Word-Dokument mit Frage und Antwort."""
    doc = DocxDocument()

    # Titel
    title = doc.add_heading("MedReg Intelligence — Analyse", level=1)
    title.runs[0].font.color.rgb = RGBColor(26, 58, 92)

    # Zeitstempel
    doc.add_paragraph(f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}", style="Intense Quote")

    # Frage
    doc.add_heading("Frage", level=2)
    doc.add_paragraph(question)

    # Antwort
    doc.add_heading("Antwort", level=2)
    for line in answer.split("\n"):
        if line.startswith("| "):
            # Tabellenzeilen als Text (Word-Tabellen aus Markdown wären zu komplex)
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(9)
            p.style.font.name = "Consolas"
        elif line.startswith("# ") or line.startswith("## ") or line.startswith("### "):
            clean = line.lstrip("# ").strip()
            doc.add_heading(clean, level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            # Fettdruck erkennen: **text**
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    if i % 2 == 1:  # ungerade = fett
                        run.bold = True
        else:
            doc.add_paragraph()  # Leerzeile

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Erstellt mit MedReg Intelligence — Lesemann AI Solutions & Consulting")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def answer_export_buttons(text, question, key):
    """Zeigt Export-Optionen für eine Antwort: Word-Download + kopierbarer Text."""
    col1, col2 = st.columns([1, 1])
    with col1:
        docx_buf = generate_answer_docx(question, text)
        st.download_button(
            label="📥 Als Word speichern",
            data=docx_buf,
            file_name=f"MedReg_Analyse_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"docx_{key}"
        )
    with col2:
        with st.popover("📋 Text kopieren"):
            st.code(text, language=None)

# --- WILLKOMMENS-BEREICH ---
if vectorstore and not st.session_state.chat_history:
    st.markdown("""
    Willkommen! Dieses Tool beantwortet Ihre Fragen zu **Medizinprodukte-Regulierung**
    auf Basis der tatsächlichen Gesetzestexte — mit exakten Quellenangaben.

    **Enthaltene Gesetze:** EU MDR, MPDG (DE), MepV (CH), UK MDR 2002

    ---
    **Beispielfragen zum Einstieg:**
    """)

    example_questions = [
        "Welche Anforderungen stellt die EU MDR an die klinische Bewertung von Medizinprodukten?",
        "Vergleiche die Klassifizierungsregeln für Medizinprodukte in der EU und UK.",
        "Was regelt das MPDG im Vergleich zur EU MDR?",
        "Welche Pflichten haben Wirtschaftsakteure nach der MepV (Schweiz)?",
    ]

    for q in example_questions:
        if st.button(f"💬 {q}", key=f"example_{hash(q)}", use_container_width=True):
            st.session_state.pending_question = q
            st.rerun()

# --- CHAT MIT RERANKING-LOGIK ---
if vectorstore:
    last_question = ""
    for i, message in enumerate(st.session_state.chat_history):
        if message.type == "human":
            last_question = message.content
        with st.chat_message(message.type):
            st.markdown(message.content)
            if message.type == "ai":
                answer_export_buttons(message.content, last_question, f"hist_{i}")

    # Eingabe: entweder aus Chat-Input oder aus Beispielfrage
    prompt = st.chat_input("Ihre regulatorische Frage...")
    if not prompt and "pending_question" in st.session_state:
        prompt = st.session_state.pending_question
        del st.session_state.pending_question

    if prompt:
        st.session_state.chat_history.append(HumanMessage(content=prompt))
        with st.chat_message("user"): st.markdown(prompt)

        with st.chat_message("assistant"):
            # 1. Retrieval: Breite Suche (20 Kandidaten)
            candidate_count = 20 if reranker else 10
            docs = vectorstore.as_retriever(search_kwargs={"k": candidate_count}).invoke(prompt)

            # 2. Reranking: Cohere wählt die besten 10 aus
            if reranker and docs:
                try:
                    rerank_response = reranker.rerank(
                        model="rerank-v3.5",
                        query=prompt,
                        documents=[d.page_content for d in docs],
                        top_n=10
                    )
                    reranked_indices = [r.index for r in rerank_response.results]
                    docs = [docs[i] for i in reranked_indices]
                except Exception as e:
                    st.warning(f"⚠️ Reranking fehlgeschlagen, nutze Standard-Ranking: {e}")

            context = "\n\n".join([
                f"QUELLE: {d.metadata.get('document_title', d.metadata.get('source_document', '?'))} "
                f"- {d.metadata.get('article_id', '')} "
                f"({d.metadata.get('jurisdiction', '?')}, S.{d.metadata.get('page', '?')})\n"
                f"{d.page_content}"
                for d in docs
            ])

            # 2. Modell & Prompt
            llm = AzureChatOpenAI(azure_deployment="gpt-4.1", api_version=API_VERSION, streaming=True, temperature=0)
            sys_msg = """Du bist ein Regulatory Intelligence Experte für Medizinprodukte.

Regeln:
1. Trenne Jurisdiktionen strikt: EU (MDR), Deutschland (MPDG), Schweiz (MepV), UK (UK MDR 2002).
2. Nutze Tabellen für Vergleiche zwischen Jurisdiktionen.
3. Belege JEDE Aussage mit der exakten Quelle: Gesetz, Artikel/§ und Jurisdiktion.
4. Wenn eine Frage mehrere Jurisdiktionen betrifft, zeige die Unterschiede klar auf.
5. Antworte in der Sprache der Frage (Deutsch oder Englisch)."""

            chain = ChatPromptTemplate.from_messages([
                ("system", sys_msg),
                MessagesPlaceholder(variable_name="chat_history"),
                ("human", "Kontext:\n{context}\n\nFrage: {input}")
            ]) | llm | StrOutputParser()

            # 3. Stream & Store
            full_res = st.write_stream(chain.stream({"input": prompt, "context": context, "chat_history": st.session_state.chat_history[:-1]}))
            answer_export_buttons(full_res, prompt, "latest")
            st.session_state.chat_history.append(AIMessage(content=full_res))
            st.rerun()
